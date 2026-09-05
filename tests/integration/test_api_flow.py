"""Full API flow on the real sample cards. Skipped unless IQAMA_SAMPLES_DIR points at sample1_individual.jpg,
sample2_establishment.jpg, sample3_company.jpg (never committed to the repo)."""
import os
import time
from pathlib import Path

import pytest

SAMPLES = Path(os.environ.get("IQAMA_SAMPLES_DIR", ""))
pytestmark = pytest.mark.skipif(not (SAMPLES / "sample1_individual.jpg").exists(), reason="sample images not available")


def _wait(api, bid, timeout=400):
    t0 = time.time()
    while time.time() - t0 < timeout:
        b = api.get(f"/api/v1/batches/{bid}").json()
        if b["status"] == "DONE":
            return b
        time.sleep(2)
    raise AssertionError("batch did not finish")


@pytest.fixture
def restore_occupations(api):
    """The flow adds an occupation through the API; put config/occupations.csv back afterwards."""
    orig = api.get("/api/v1/rules/files/occupations.csv").json()["content"]
    yield
    api.put("/api/v1/rules/files/occupations.csv", json={"content": orig})


def test_full_flow(api, restore_occupations):
    os.environ["IQAMA_OCR_PROVIDER"] = os.environ.get("IQAMA_TEST_OCR_PROVIDER", "easyocr")
    from app.api.deps import batch_service
    batch_service().provider_name = os.environ["IQAMA_OCR_PROVIDER"]

    bid = api.post("/api/v1/batches", json={"name": "samples"}).json()["id"]
    files = [("files", (n, (SAMPLES / n).read_bytes(), "image/jpeg")) for n in
             ["sample1_individual.jpg", "sample2_establishment.jpg", "sample3_company.jpg"]]
    files.append(("files", ("bad.txt", b"not an image", "text/plain")))
    up = api.post(f"/api/v1/batches/{bid}/documents", files=files).json()
    assert len(up["accepted"]) == 3 and len(up["rejected"]) == 1
    assert api.post(f"/api/v1/batches/{bid}/process").status_code == 202
    b = _wait(api, bid)
    assert b["processed"] == 3 and b["summary"]["total"] == 3

    docs = api.get(f"/api/v1/batches/{bid}/documents").json()
    by_name = {d["filename"]: d for d in docs}
    s1, s2, s3 = by_name["sample1_individual.jpg"], by_name["sample2_establishment.jpg"], by_name["sample3_company.jpg"]
    for d in (s1, s2, s3):
        assert d["status"] == "DONE", d["error"]
        assert d["decision"]["status"] in ("REJECTED", "MANUAL_REVIEW")       # never APPROVED without a human
        assert d["fields"]["iqama_no"]["normalized"].endswith("******")        # masked by default
    # deterministic facts we know the OCR gets right at this resolution
    assert s2["fields"]["employer_id"]["normalized"].startswith("7034")
    assert s3["fields"]["employer_id"]["normalized"].startswith("7015")
    assert s1["fields"]["employer_id"]["normalized"].startswith("1052")
    assert s2["fields"]["name_en"]["normalized"] == "ABUBAKER ABASS ALZEBIR ABASS"
    assert s3["fields"]["expiry_date"]["normalized"] == "2026-04-02"
    assert s1["fields"]["expiry_date"]["normalized"] == "2022-12-11"

    # unmasked read is audited
    full = api.get(f"/api/v1/documents/{s2['id']}", params={"unmask": "true"}, headers={"X-Actor": "reviewer1"}).json()
    assert full["fields"]["iqama_no"]["normalized"] == "2627946219"
    assert any(a["action"] == "PII_UNMASKED" for a in api.get("/api/v1/audit").json())

    # image is available for review while MANUAL_REVIEW
    assert api.get(f"/api/v1/documents/{s2['id']}/image").status_code == 200

    # individual employer (prefix 1, checksum ok) rejects directly even before review (Rules v1 D5)
    assert s1["decision"]["status"] == "REJECTED" and "Individual Employer" in s1["decision"]["reasons"]
    # reviewer confirms sample 1's occupation + expiry -> decision re-runs and lists all 3 reasons
    r = api.patch(f"/api/v1/documents/{s1['id']}/fields", json={"fields": {"occupation": "سائق خاص", "expiry_date": "2022-12-11"}},
                  headers={"X-Actor": "reviewer1"}).json()
    assert r["decision"]["status"] == "REJECTED" and len(r["decision"]["reasons"]) == 3, r["decision"]
    assert r["decision"]["version"] == 2

    # reviewer approves sample 2 after adding its occupation as eligible and confirming the low-confidence reads
    api.post("/api/v1/rules/occupations", json={"occupation_ar": "عامل تحميل وتنزيل", "occupation_en": "Loading Worker", "eligible": True})
    r = api.patch(f"/api/v1/documents/{s2['id']}/fields", json={"fields": {"expiry_date": "2026-12-30", "occupation": "عامل تحميل وتنزيل"}},
                  headers={"X-Actor": "reviewer1"}).json()
    assert r["decision"]["status"] == "MANUAL_REVIEW" and r["decision"]["recommendation"] == "RECOMMEND_APPROVE", r["decision"]
    r = api.post(f"/api/v1/documents/{s2['id']}/review", json={"status": "APPROVED", "note": "ok"}, headers={"X-Actor": "reviewer1"}).json()
    assert r["decision"]["status"] == "APPROVED" and r["decision"]["is_final"]
    assert api.get(f"/api/v1/documents/{s2['id']}/image").status_code == 200   # kept until the permit request is generated

    # permit request from the customer's Word template (only when the real template is available locally)
    tpl = os.environ.get("IQAMA_PERMIT_TEMPLATE")
    if tpl and Path(tpl).exists():
        up = api.put("/api/v1/templates/permit", files={"file": ("t.docx", Path(tpl).read_bytes(), "application/vnd.openxmlformats-officedocument.wordprocessingml.document")})
        assert up.status_code == 200, up.text
        api.patch(f"/api/v1/batches/{bid}/project", json={"requesting_companies": ["قدرة العربية"], "project": {"name": "مشروع اختبار", "location": "الرياض", "work_start": "2026/09/10", "work_end_expected": "2026/12/31"}})
        doc2 = api.get(f"/api/v1/documents/{s2['id']}").json()
        assert doc2["company_source"] in ("CARD", "AUTO", "MANUAL") and doc2["company_options"]
        api.patch(f"/api/v1/documents/{s2['id']}/company", json={"company_name": "قدرة العربية"})
        p = api.get(f"/api/v1/batches/{bid}/permit", params={"format": "docx"})
        assert p.status_code == 200 and p.content[:2] == b"PK", p.text
        from docx import Document
        import io as _io
        d = Document(_io.BytesIO(p.content))
        team = d.tables[1]
        assert team.rows[1].cells[4].text == "ابوبكر عباس الزبير عباس" and team.rows[1].cells[2].text == "قدرة العربية"
        assert team.rows[1].cells[1].text.replace("\u200e", "") == "2627946219"   # the form needs the full number
        assert d.tables[2].rows and "ابوبكر" in d.tables[2].rows[0].cells[1].text   # card image + caption
        assert api.get(f"/api/v1/batches/{bid}").json()["permit_exported_at"]

    # exports
    x = api.get(f"/api/v1/batches/{bid}/export", params={"format": "xlsx"})
    assert x.status_code == 200 and len(x.content) > 5000
    c = api.get(f"/api/v1/batches/{bid}/export", params={"format": "csv"})
    assert c.status_code == 200 and "Final Decision" in c.text and "2627946219" not in c.text
    summ = api.get(f"/api/v1/batches/{bid}").json()["summary"]
    assert summ["approved"] == 1 and summ["rejected"] >= 1
    # queue no longer contains the approved doc
    assert s2["id"] not in [q["document_id"] for q in api.get("/api/v1/review/queue").json()]
