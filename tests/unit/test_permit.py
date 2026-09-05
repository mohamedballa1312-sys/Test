"""Permit generator on a synthetic template with the same 5-table structure as the customer's form (no letterhead/stamp)."""
import io
import zipfile
from datetime import date

from PIL import Image

from app.services.permit import PermitData, PermitGenerator, Worker, hijri_str

W = 'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships" xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing" xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main" xmlns:pic="http://schemas.openxmlformats.org/drawingml/2006/picture"'


def _p(text=""):
    return f'<w:p><w:r><w:rPr><w:sz w:val="22"/></w:rPr><w:t xml:space="preserve">{text}</w:t></w:r></w:p>' if text else "<w:p/>"


def _tc(text=""):
    return f'<w:tc><w:tcPr><w:tcW w:w="2000" w:type="dxa"/></w:tcPr>{_p(text)}</w:tc>'


def _stamp_run(rid="rId9", docpr=7):
    return (f'<w:r><w:drawing><wp:inline><wp:extent cx="700000" cy="380000"/><wp:docPr id="{docpr}" name="stamp"/>'
            f'<a:graphic><a:graphicData uri="http://schemas.openxmlformats.org/drawingml/2006/picture"><pic:pic><pic:nvPicPr><pic:cNvPr id="0" name="s"/><pic:cNvPicPr/></pic:nvPicPr>'
            f'<pic:blipFill><a:blip r:embed="{rid}"/></pic:blipFill><pic:spPr><a:xfrm><a:off x="0" y="0"/><a:ext cx="700000" cy="380000"/></a:xfrm></pic:spPr></pic:pic></a:graphicData></a:graphic></wp:inline></w:drawing></w:r>')


def synthetic_template() -> bytes:
    t0 = "<w:tbl><w:tblGrid><w:gridCol w:w=\"4000\"/><w:gridCol w:w=\"2000\"/></w:tblGrid>" + "".join(
        f"<w:tr>{_tc('(يتم تزويده من المستخدم)')}{_tc(lbl)}</w:tr>" for lbl in ["اسم المشروع", "الموقع", "تاريخ بداية الأعمال", "التاريخ المتوقع"]) + "</w:tbl>"
    header = "<w:tr>" + "".join(_tc(h) for h in ["ملاحظات", "رقم الهوية / الاقامة", "اسم الشركة", "الجنسية", "الاسم", "م"]) + "</w:tr>"
    row1 = ("<w:tr>" + f'<w:tc><w:tcPr><w:tcW w:w="2000" w:type="dxa"/></w:tcPr><w:p><w:r><w:rPr><w:sz w:val="22"/></w:rPr><w:t>عند الموافقه استخدم الختم</w:t></w:r>{_stamp_run()}</w:p></w:tc>'
            + "".join(_tc(x) for x in ["ييم استخراجه من النظام", "(ييم استخراجه من النظام)", "(ييم استخراجه من النظام)", "(ييم استخراجه من النظام)", "1"]) + "</w:tr>")
    def empty_row(n): return "<w:tr>" + "".join(_tc("") for _ in range(5)) + _tc(str(n)) + "</w:tr>"
    t1 = "<w:tbl><w:tblGrid>" + "<w:gridCol w:w=\"2000\"/>" * 6 + "</w:tblGrid>" + header + row1 + "".join(empty_row(n) for n in range(2, 13)) + "</w:tbl>"
    t2 = "<w:tbl><w:tblGrid>" + "<w:gridCol w:w=\"2000\"/>" * 6 + "</w:tblGrid>" + "".join(empty_row(n) for n in range(13, 28)) + "</w:tbl>"
    t3 = "<w:tbl><w:tblGrid><w:gridCol w:w=\"4500\"/><w:gridCol w:w=\"4500\"/></w:tblGrid><w:tr>" + f'<w:tc><w:tcPr><w:tcW w:w="4500" w:type="dxa"/></w:tcPr><w:p>{_stamp_run("rId9", 8)}</w:p></w:tc>' + _tc("(صوره من الاقامه)") + "</w:tr></w:tbl>"
    t4 = "<w:tbl><w:tblGrid><w:gridCol w:w=\"4500\"/><w:gridCol w:w=\"4500\"/></w:tblGrid>" + "".join("<w:tr>" + _tc("") + _tc("") + "</w:tr>" for _ in range(3)) + "</w:tbl>"
    doc = (f'<?xml version="1.0" encoding="UTF-8" standalone="yes"?><w:document {W}><w:body>'
           f'{_p("التـاريخ: ) تاريخ الاصدار)   الموافق: (تاريخ الاصدار هجري)")}{_p("طلب تصريح")}{t0}{_p("فريق العمل")}{t1}{t2}{_p("الهوية الوطنية / الاقامة")}{t3}{t4}'
           f'<w:sectPr/></w:body></w:document>')
    stamp = io.BytesIO(); Image.new("RGB", (60, 30), "blue").save(stamp, "PNG")
    rels = ('<?xml version="1.0" encoding="UTF-8"?><Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
            '<Relationship Id="rId9" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/image" Target="media/stamp.png"/></Relationships>')
    ct = ('<?xml version="1.0" encoding="UTF-8"?><Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">'
          '<Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/><Default Extension="xml" ContentType="application/xml"/>'
          '<Default Extension="png" ContentType="image/png"/><Default Extension="jpg" ContentType="image/jpeg"/>'
          '<Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/></Types>')
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w") as z:
        z.writestr("[Content_Types].xml", ct); z.writestr("_rels/.rels", '<?xml version="1.0"?><Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships"><Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="word/document.xml"/></Relationships>')
        z.writestr("word/document.xml", doc); z.writestr("word/_rels/document.xml.rels", rels); z.writestr("word/media/stamp.png", stamp.getvalue())
    return buf.getvalue()


def _card_png():
    b = io.BytesIO(); Image.new("RGB", (400, 250), "white").save(b, "PNG"); return b.getvalue()


def _read(docx: bytes):
    from docx import Document
    return Document(io.BytesIO(docx))


def test_fills_dates_project_team_and_images():
    gen = PermitGenerator(synthetic_template())
    workers = [Worker(name=f"عامل {i}", nationality="السودان", company="قدرة العربية", id_number=f"24000000{i:02d}", approved=True, image_png=_card_png()) for i in range(1, 4)]
    workers.append(Worker(name="مرفوض", nationality="مصر", company="x", id_number="2400000099", approved=False))
    data = PermitData(issue_date=date(2026, 9, 5), project_name="مشروع أ", project_location="الرياض", work_start="2026/09/10", work_end_expected="2026/12/31", workers=workers)
    out = gen.render_docx(data)
    d = _read(out)
    text = "\n".join(p.text for p in d.paragraphs)
    assert "2026/09/05" in text and hijri_str(date(2026, 9, 5)) in text
    assert d.tables[0].rows[0].cells[0].text == "مشروع أ" and d.tables[0].rows[1].cells[0].text == "الرياض"
    team = d.tables[1]
    assert [c.text for c in team.rows[1].cells][1:] == ["2400000001", "قدرة العربية", "السودان", "عامل 1", "1"]
    assert team.rows[4].cells[4].text == "مرفوض" and team.rows[5].cells[4].text == ""     # unapproved still listed (caller filters), rows beyond stay empty
    xml = zipfile.ZipFile(io.BytesIO(out)).read("word/document.xml").decode()
    assert xml.count('r:embed="rId9"') >= 3 + 3           # stamp in 3 approved notes cells + on 3 images (template stamps removed/replaced)
    assert len(zipfile.ZipFile(io.BytesIO(out)).namelist()) >= 6 and any("iqama_" in n for n in zipfile.ZipFile(io.BytesIO(out)).namelist())
    # images table: 3 images -> 2 rows, captions present
    imgs = d.tables[3]
    assert len(imgs.rows) == 2 and "عامل 1" in imgs.rows[0].cells[1].text and "عامل 3" in imgs.rows[1].cells[1].text


def test_more_than_27_workers_adds_rows():
    gen = PermitGenerator(synthetic_template())
    workers = [Worker(name=f"w{i}", nationality="n", company="c", id_number=str(2400000000 + i), approved=True) for i in range(1, 31)]
    out = gen.render_docx(PermitData(issue_date=date(2026, 1, 1), workers=workers))
    d = _read(out)
    assert len(d.tables[2].rows) == 18 and d.tables[2].rows[-1].cells[5].text == "30" and d.tables[2].rows[-1].cells[4].text == "w30"


def test_rejects_wrong_template():
    import pytest
    from app.services.permit import PermitTemplateError
    with pytest.raises(PermitTemplateError):
        PermitGenerator(b"not a docx")
